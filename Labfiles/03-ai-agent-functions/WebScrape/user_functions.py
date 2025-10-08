import json
from typing import Any, Callable, Set

# For web scraping
import requests
from bs4 import BeautifulSoup

# Function to scrape info from a URL
def scrape_url_info(url: str) -> str:
    """
    Scrapes all visible text content from the given URL.
    Returns a JSON string with the URL and the extracted text or error message.
    """
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        # Remove script and style elements
        for tag in soup(['script', 'style', 'noscript']):
            tag.decompose()
        # Get visible text
        text = soup.get_text(separator=' ', strip=True)
        # Optionally, limit the length to avoid huge outputs
        max_length = 5000
        if len(text) > max_length:
            text = text[:max_length] + '... [truncated]'
        result = {
            'url': url,
            'text': text
        }
        return json.dumps(result)
    except Exception as e:
        return json.dumps({'error': str(e), 'url': url})




# Function to save output to a Word document
def save_output_to_word(output: str, filename: str = "output.docx") -> str:
    """
    Saves the given output (string or JSON string) to a Word document.
    If output is a JSON string with 'url' and 'text', formats them nicely.
    Returns the filename used.
    """
    try:
        from docx import Document
    except ImportError:
        return "python-docx package is not installed. Please install it to use this function."

    doc = Document()
    try:
        # Try to parse as JSON
        data = json.loads(output)
        if isinstance(data, dict) and 'url' in data and 'text' in data:
            doc.add_heading(f"Scraped Content from {data['url']}", level=1)
            doc.add_paragraph(data['text'])
        else:
            doc.add_paragraph(output)
    except Exception:
        # Not JSON, just add as plain text
        doc.add_paragraph(output)
    doc.save(filename)
    return filename

# Define a set of callable functions
user_functions: Set[Callable[..., Any]] = {
    scrape_url_info,
    save_output_to_word
}


